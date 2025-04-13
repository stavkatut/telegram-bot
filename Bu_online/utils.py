import os
import aiohttp
import asyncio
import json
import hashlib
from datetime import datetime, timedelta
from docx import Document
from dotenv import load_dotenv
import pandas as pd
from pathlib import Path
import logging
from logging.handlers import RotatingFileHandler
from aiohttp import ClientTimeout, TCPConnector
from typing import Optional, Dict, Any

# Настройка логирования
def setup_logging():
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)

    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    # Логирование в файл с ротацией
    file_handler = RotatingFileHandler(
        'accountant.log',
        maxBytes=5*1024*1024,  # 5 MB
        backupCount=3,
        encoding='utf-8'
    )
    file_handler.setFormatter(formatter)

    # Логирование в консоль
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)

    logger.addHandler(file_handler)
    logger.addHandler(console_handler)

setup_logging()
logger = logging.getLogger(__name__)

load_dotenv()

class AIAccountantCore:
    """Полнофункциональное ядро бухгалтерского помощника с DeepSeek API"""
    
    def __init__(self):
        self.api_url = "https://api.deepseek.com/v1/chat/completions"
        self.api_key = os.getenv("DEEPSEEK_API_KEY")
        self.session = None
        self.timeout = ClientTimeout(total=30, connect=10)  # Увеличенные таймауты
        self.retries = 3
        self.backoff_base = 2
        
        # База знаний по налогам
        self.tax_knowledge = {
            "УСН": {
                "rate": 0.06,
                "deadline": "25 число следующего месяца",
                "forms": ["КНИД", "Декларация УСН"]
            },
            "УСН15": {
                "rate": 0.15,
                "deadline": "25 число следующего месяца",
                "forms": ["КНИД", "Декларация УСН"]
            },
            "НДФЛ": {
                "rate": 0.13,
                "deadline": "15 июля следующего года",
                "forms": ["3-НДФЛ"]
            }
        }
        
        # Локальная база знаний для fallback
        self.local_knowledge = {
            "усн": {
                "response": (
                    "📌 УСН (упрощенная система налогообложения):\n"
                    "• Ставка: 6% от доходов или 15% от (доходы - расходы)\n"
                    "• Отчетность: Декларация УСН (до 30 апреля)\n"
                    "• Уплата: авансовые платежи до 25 числа\n"
                    "• НПА: НК РФ ст. 346.12-346.27"
                ),
                "keywords": ["усн", "упрощен"]
            },
            "ндфл": {
                "response": (
                    "📌 НДФЛ (налог на доходы физлиц):\n"
                    "• Ставка: 13% (резиденты), 30% (нерезиденты)\n"
                    "• Срок уплаты: не позднее 15 июля\n"
                    "• Форма: 3-НДФЛ\n"
                    "• НПА: НК РФ ст. 207-233"
                ),
                "keywords": ["ндфл", "подоходный"]
            }
        }

    async def ensure_session(self):
        """Инициализация и поддержание сессии"""
        if self.session is None or self.session.closed:
            connector = TCPConnector(
                limit=10,
                force_close=False,
                enable_cleanup_closed=True,
                ssl=False  # Отключить SSL верификацию при проблемах
            )
            self.session = aiohttp.ClientSession(
                connector=connector,
                timeout=self.timeout,
                headers={
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json",
                    "User-Agent": "AccountingBot/3.0",
                    "Accept": "application/json"
                }
            )

    async def ask_ai(self, prompt: str) -> str:
        """Улучшенный запрос к DeepSeek API с обработкой ошибок"""
        await self.ensure_session()
        
        # Проверка локальной базы знаний
        local_response = self._get_local_response(prompt)
        if local_response:
            return local_response

        payload = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": "Ты экспертный бухгалтер. Отвечай кратко, но информативно. Формат: 1) Суть 2) Нормативная база 3) Рекомендации"
                },
                {
                    "role": "user",
                    "content": prompt[:2000]  # Лимит длины
                }
            ],
            "temperature": 0.3,
            "max_tokens": 1000,
            "stream": False
        }

        last_error = None
        for attempt in range(self.retries):
            try:
                async with self.session.post(
                    self.api_url,
                    json=payload,
                    timeout=self.timeout
                ) as response:
                    # Логирование сырого ответа
                    raw_response = await response.text()
                    logger.debug(f"API Response: {raw_response}")

                    if response.status != 200:
                        error_data = await response.json()
                        raise aiohttp.ClientError(
                            f"API Error {response.status}: {error_data.get('error', {}).get('message', 'Unknown error')}"
                        )

                    data = await response.json()
                    return data["choices"][0]["message"]["content"]

            except Exception as e:
                last_error = str(e)
                logger.error(f"Attempt {attempt+1} failed: {last_error}")
                if attempt == self.retries - 1:
                    return self._get_local_response(prompt) or f"⚠️ Ошибка сервиса: {last_error}"
                await asyncio.sleep(self.backoff_base ** attempt)
        
        return self._get_local_response(prompt) or "⚠️ Ошибка обработки запроса"

    def _get_local_response(self, prompt: str) -> Optional[str]:
        """Поиск в локальной базе знаний"""
        prompt_lower = prompt.lower()
        for topic in self.local_knowledge.values():
            if any(kw in prompt_lower for kw in topic["keywords"]):
                return topic["response"]
        return None

    async def check_connection(self) -> bool:
        """Проверка доступности API"""
        try:
            await self.ensure_session()
            async with self.session.get(
                "https://api.deepseek.com/v1/models",
                timeout=ClientTimeout(total=5)
            ) as response:
                return response.status == 200
        except Exception as e:
            logger.error(f"Connection check failed: {str(e)}")
            return False

    def calculate_tax(self, income: float, system: str = "УСН", region: str = None) -> dict:
        """Расширенный расчет налогов"""
        try:
            income = float(income)
            if system not in self.tax_knowledge:
                raise ValueError("Неподдерживаемая система налогообложения")
                
            tax_info = self.tax_knowledge[system]
            tax = income * tax_info["rate"]
            
            result = {
                "system": system,
                "income": income,
                "rate": tax_info["rate"],
                "tax": tax,
                "deadline": tax_info["deadline"],
                "forms": tax_info.get("forms", [])
            }
            
            if region:
                result["notes"] = "Учтите региональные коэффициенты"
                
            return result
            
        except Exception as e:
            logger.error(f"Tax calculation error: {str(e)}")
            raise

    async def generate_document(self, doc_type: str, **kwargs) -> str:
        """Генератор документов с базовой валидацией"""
        try:
            if doc_type == "Договор":
                return await self._generate_contract(**kwargs)
            elif doc_type == "Акт":
                return await self._generate_act(**kwargs)
            else:
                raise ValueError("Неизвестный тип документа")
        except Exception as e:
            logger.error(f"Document generation failed: {str(e)}")
            raise

    async def _generate_contract(self, client: str, amount: float, service: str) -> str:
        """Генерация договора"""
        doc = Document()
        
        # Шапка документа
        doc.add_heading(f'ДОГОВОР № {datetime.now().strftime("%Y%m%d")}', 0)
        doc.add_paragraph(f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_paragraph(f"Исполнитель: ООО 'БухПрофи'")
        doc.add_paragraph(f"Заказчик: {client}")
        
        # Условия договора
        doc.add_heading('1. Предмет договора', level=1)
        doc.add_paragraph(f"Исполнитель обязуется оказать услуги: {service}")
        
        doc.add_heading('2. Стоимость услуг', level=1)
        doc.add_paragraph(f"Общая стоимость услуг: {amount:,.2f} руб. (НДС не облагается)")
        
        # Сохранение
        os.makedirs("documents", exist_ok=True)
        filename = f"documents/Договор_{client}_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(filename)
        
        return filename

    def analyze_excel(self, file_path: str) -> dict:
        """Анализ финансовых данных"""
        try:
            df = pd.read_excel(file_path)
            
            # Проверка обязательных колонок
            required_cols = ['Дата', 'Сумма']
            missing = [col for col in required_cols if col not in df.columns]
            if missing:
                raise ValueError(f"Отсутствуют колонки: {', '.join(missing)}")
                
            df['Дата'] = pd.to_datetime(df['Дата'], errors='coerce')
            
            return {
                "period": {
                    "start": df['Дата'].min().strftime('%d.%m.%Y'),
                    "end": df['Дата'].max().strftime('%d.%m.%Y')
                },
                "total_income": df[df['Сумма'] > 0]['Сумма'].sum(),
                "total_expenses": abs(df[df['Сумма'] < 0]['Сумма'].sum()),
                "tax_optimization": self._get_tax_optimization(df)
            }
        except Exception as e:
            logger.error(f"Excel analysis error: {str(e)}")
            raise

    def _get_tax_optimization(self, df: pd.DataFrame) -> str:
        """Рекомендации по оптимизации налогов"""
        profit = df['Сумма'].sum()
        if profit > 3000000:
            return "Рекомендуем ОСНО с НДС"
        elif profit > 1000000:
            return "Оптимально УСН 15% (доходы минус расходы)"
        else:
            return "Можно применять УСН 6% или патент"

    async def close(self):
        """Корректное завершение сессии"""
        if self.session and not self.session.closed:
            await self.session.close()